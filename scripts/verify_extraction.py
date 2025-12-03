import requests
from docx import Document

# 1. Create a dummy DOCX
doc = Document()
doc.add_heading('CaseStar Test Document', 0)
doc.add_paragraph('This is a test document to verify extraction.')
doc.add_paragraph('It contains multiple paragraphs to check formatting.')
doc.add_paragraph('Entity: John Doe (Plaintiff) vs. Acme Corp (Defendant).')
doc.save('verify_extraction.docx')

# 2. Upload to API
url = 'http://localhost:50000/api/upload'
files = {'file': open('verify_extraction.docx', 'rb')}

try:
    print("Uploading verify_extraction.docx...")
    response = requests.post(url, files=files)

    if response.status_code == 200:
        data = response.json()
        print("\n✅ Upload Successful!")
        print(f"Filename: {data['filename']}")
        print(f"Extracted Text Preview:\n{'-'*20}")
        print(data['extracted_text'])
        print(f"{'-'*20}")

        if "John Doe" in data['extracted_text']:
            print("✅ Extraction Verified: Found 'John Doe'")
        else:
            print("❌ Extraction Failed: Content missing")

    else:
        print(f"❌ Upload Failed: {response.status_code} - {response.text}")

except Exception as e:
    print(f"❌ Error: {e}")
