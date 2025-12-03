import fitz  # PyMuPDF
from docx import Document
import io
import logging

logger = logging.getLogger(__name__)


def extract_text_from_pdf(file_content: bytes) -> str:
    """Extract text from PDF content using PyMuPDF"""
    try:
        text = ""
        with fitz.open(stream=file_content, filetype="pdf") as doc:

            # Check for encryption
            if doc.is_encrypted:
                # Try to authenticate with empty password
                doc.authenticate("")
                if doc.is_encrypted:
                    logger.warning("PDF is encrypted and cannot be read")
                    return "Error: PDF is password protected"

            for page in doc:
                text += page.get_text()

        return text
    except Exception as e:
        logger.error(f"PDF extraction error: {e}")
        return f"Error extracting text from PDF: {str(e)}"


def extract_text_from_docx(file_content: bytes) -> str:
    """Extract text from DOCX content using python-docx"""
    try:
        doc = Document(io.BytesIO(file_content))
        full_text = []

        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)

        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [
                    cell.text for cell in row.cells if cell.text.strip()
                ]
                if row_text:
                    full_text.append(" | ".join(row_text))

        return "\n\n".join(full_text)
    except Exception as e:
        logger.error(f"DOCX extraction error: {e}")
        return f"Error extracting text from DOCX: {str(e)}"


async def process_document_content(filename: str, content: bytes) -> str:
    """Route document processing based on file extension"""
    filename = filename.lower()

    if filename.endswith('.pdf'):
        return extract_text_from_pdf(content)
    elif filename.endswith('.docx'):
        return extract_text_from_docx(content)
    elif filename.endswith('.txt'):
        try:
            return content.decode('utf-8')
        except UnicodeDecodeError:
            return content.decode('utf-8', errors='ignore')
    else:
        return "Unsupported file format"
